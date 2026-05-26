"""
Enhanced Document Processor - Multi-format Support
Handles PDF, Word, Markdown with tables and images
"""
import logging
import os
from pathlib import Path

logger = logging.getLogger(__name__)


class EnhancedDocumentProcessor:
    """
    Supports:
    - Markdown (.md) current
    - PDF with text (.pdf)
    - PDF with tables (.pdf)
    - PDF with images (.pdf) via Gemini Vision
    - Word documents (.docx)
    - Scanned PDFs via OCR
    """

    def __init__(self, gemini_api_key: str = None):
        self.api_key = gemini_api_key

    def extract_from_pdf(self, file_path: str) -> str:
        """Extract text, tables and images from PDF."""
        try:
            import fitz
            doc = fitz.open(file_path)
            full_text = []

            for page_num, page in enumerate(doc):
                text = page.get_text()
                if text.strip():
                    full_text.append(f"[Page {page_num+1}]\n{text}")

                tables = page.find_tables()
                for table in tables:
                    df = table.to_pandas()
                    table_text = df.to_string(index=False)
                    full_text.append(f"[Table on Page {page_num+1}]\n{table_text}")

                if self.api_key:
                    for img_idx, img in enumerate(page.get_images()):
                        try:
                            img_text = self._describe_image_with_gemini(doc, img, page_num, img_idx)
                            if img_text:
                                full_text.append(f"[Image on Page {page_num+1}]\n{img_text}")
                        except Exception as e:
                            logger.warning(f"Image extraction failed: {e}")

            return "\n\n".join(full_text)

        except ImportError:
            logger.error("PyMuPDF not installed! pip install pymupdf")
            return ""
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return ""

    def _describe_image_with_gemini(self, doc, img_ref, page_num, img_idx) -> str:
        """Use Gemini Vision to describe images in documents."""
        try:
            from google import genai
            import base64

            xref = img_ref[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            image_b64 = base64.b64encode(image_bytes).decode()

            client = genai.Client(api_key=self.api_key)
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=[{
                    "parts": [
                        {"inline_data": {"mime_type": f"image/{image_ext}", "data": image_b64}},
                        {"text": "This is an image from an HR policy document. Describe what this image shows in detail, especially any text, charts, tables, or organizational information visible."}
                    ]
                }]
            )
            return response.text

        except Exception as e:
            logger.warning(f"Gemini Vision failed: {e}")
            return ""

    def extract_from_docx(self, file_path: str) -> str:
        """Extract text and tables from Word document."""
        try:
            from docx import Document
            doc = Document(file_path)
            full_text = []

            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            for table_idx, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                full_text.append(f"[Table {table_idx+1}]\n" + "\n".join(rows))

            return "\n\n".join(full_text)

        except ImportError:
            logger.error("python-docx not installed! pip install python-docx")
            return ""


    def extract_from_excel(self, file_path: str) -> str:
        """Extract text and data from Excel files."""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            full_text = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                full_text.append(f"[Sheet: {sheet_name}]")
                rows = []
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(c.strip() for c in cells):
                        rows.append(" | ".join(cells))
                if rows:
                    full_text.append("\n".join(rows))
            return "\n\n".join(full_text)
        except ImportError:
            logger.error("openpyxl not installed!")
            return ""
        except Exception as e:
            logger.error(f"Excel extraction failed: {e}")
            return ""

    def process_file(self, file_path: str) -> str:
        """Auto-detect format and extract text."""
        ext = Path(file_path).suffix.lower()
        logger.info(f"Processing {ext} file")
        if ext == ".pdf":
            return self.extract_from_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return self.extract_from_docx(file_path)
        elif ext in [".xlsx", ".xls"]:
            return self.extract_from_excel(file_path)
        elif ext in [".md", ".txt"]:
            with open(file_path, "r") as f:
                return f.read()
        else:
            logger.warning(f"Unsupported format: {ext}")
            return ""

